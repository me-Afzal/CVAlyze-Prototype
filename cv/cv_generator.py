from docx import Document
from fpdf import FPDF
import random
import os
import textwrap


names = [
    "John Doe", "Jane Smith", "Alice Johnson", "Bob Brown", "Charlie Davis",
    "Eva Wilson", "Frank Miller", "Grace Lee", "Henry Clark", "Ivy Walker"
]

locations = [
    "New York, USA", "Bangalore, India", "Delhi, India", "San Francisco, USA", "Mumbai, India",
    "London, UK", "Toronto, Canada", "Sydney, Australia", "Berlin, Germany", "Singapore"
]

linkedin_template = "https://linkedin.com/in/{}"
github_template = "https://github.com/{}"

skills_list = [
    "Python, SQL, Machine Learning, Data Analysis",
    "R, Tableau, Power BI, Statistics",
    "Deep Learning, TensorFlow, PyTorch",
    "NLP, Time Series Forecasting, Scikit-learn",
    "Cloud Computing: AWS, GCP, Azure"
]

projects_list = [
    "Retail Analytics Dashboard - Sales forecasting, pricing, recommendations",
    "Fraud Detection System - Credit card fraud detection with Random Forest",
    "E-commerce Chatbot - AI-powered chatbot using FastAPI + GPT",
    "Customer Segmentation - K-Means clustering for personalized marketing",
    "Dynamic Pricing Engine - Optimizing product prices based on demand"
]

certifications_list = [
    "AWS Certified Machine Learning - Specialty",
    "Google Cloud Professional Data Engineer",
    "Advanced SQL for Data Science - Coursera",
    "Tableau Desktop Specialist",
    "Microsoft Azure AI Fundamentals"
]

achievements_list = [
    "Won Top 10 AI Innovation Award at Walmart Hackathon 2023",
    "Published research paper on AI-driven Pricing Optimization at IEEE ICMLA 2022",
    "Recognized as Employee of the Year 2021 at Company ABC",
    "Developed award-winning ML model for customer churn reduction",
    "Presented at Data Science Summit 2022"
]

education_list = [
    "M.Tech, Data Science - Indian Institute of Technology (IIT), 2020",
    "B.E., Computer Science - Visvesvaraya Technological University (VTU), 2018",
    "M.Sc, Statistics - University of California, 2019",
    "B.Sc, Computer Science - New York University, 2017",
    "MBA, Business Analytics - London Business School, 2021"
]

# Helper functions
def pick_random(items, n=3):
    return random.sample(items, min(n, len(items)))

def clean_text(text):
    # Replace problematic characters and ensure ASCII compatibility
    return text.replace("–", "-").replace("•", "-").encode('ascii', 'ignore').decode('ascii')

def wrap_text(text, width=60):
    # Wrap long text for PDF with shorter width
    return "\n".join(textwrap.wrap(text, width=width))

# Create directory if it doesn't exist
os.makedirs("dummy_cvs", exist_ok=True)

# Generate 5 DOCX CVs
for i in range(5):
    doc = Document()
    name = names[i]
    location = locations[i]
    doc.add_heading(name, 0)
    doc.add_paragraph(f"Location: {location}")
    doc.add_paragraph(f"Email: {name.split()[0].lower()}@email.com | Phone: +91-900000000{i}")
    doc.add_paragraph(f"LinkedIn: {linkedin_template.format(name.split()[0].lower())}")
    doc.add_paragraph(f"GitHub: {github_template.format(name.split()[0].lower())}")
    doc.add_paragraph("Profile Summary: Experienced professional in data science and analytics.")
    doc.add_paragraph("Skills: " + ", ".join(pick_random(skills_list, 3)))
    doc.add_paragraph("Education: " + ", ".join(pick_random(education_list, 2)))
    doc.add_paragraph("Projects: " + ", ".join(pick_random(projects_list, 3)))
    doc.add_paragraph("Certifications: " + ", ".join(pick_random(certifications_list, 2)))
    doc.add_paragraph("Achievements: " + ", ".join(pick_random(achievements_list, 2)))
    doc.add_paragraph("Experience: Worked at Company ABC as Data Scientist")
    
    doc.save(f"dummy_cv_{i+1}.docx")

# Generate 5 PDF CVs using FPDF
for i in range(5, 10):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_margins(20, 20, 20)  # Set proper margins
    
    name = names[i]
    location = locations[i]
    
    # Header with name
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, clean_text(name), new_x="LMARGIN", new_y="NEXT", align='C')
    pdf.ln(5)
    
    # Content sections
    pdf.set_font("Helvetica", "", 10)
    
    sections = [
        f"Location: {location}",
        f"Email: {name.split()[0].lower()}@email.com | Phone: +91-900000000{i}",
        f"LinkedIn: {linkedin_template.format(name.split()[0].lower())}",
        f"GitHub: {github_template.format(name.split()[0].lower())}",
        "Profile Summary: Experienced professional in data science and analytics.",
        "Skills: " + ", ".join(pick_random(skills_list, 3)),
        "Education: " + ", ".join(pick_random(education_list, 2)),
        "Projects: " + ", ".join(pick_random(projects_list, 3)),
        "Certifications: " + ", ".join(pick_random(certifications_list, 2)),
        "Achievements: " + ", ".join(pick_random(achievements_list, 2)),
        "Experience: Worked at Company ABC as Data Scientist"
    ]
    
    for section in sections:
        clean_section = clean_text(section)
        # Split long lines to fit within page width
        lines = textwrap.wrap(clean_section, width=80)
        for line in lines:
            pdf.cell(0, 6, line, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)  # Small gap between sections
    
    pdf.output(f"dummy_cv_{i+1}.pdf")

print("10 enhanced dummy CVs generated in 'dummy_cvs' folder (5 DOCX, 5 PDF).")